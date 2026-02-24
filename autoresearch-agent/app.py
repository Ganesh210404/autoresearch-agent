import json
import io
import re
import streamlit as st
from main_langchain import run_agent

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except ImportError:
    letter = None
    canvas = None

st.set_page_config(page_title="AutoResearch Agent", page_icon="⚡", layout="wide")

# ---------- Styles ----------
st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@400;500;600;700&display=swap');

html, body, [class*="css"] {
    font-family: "Space Grotesk", sans-serif;
}

.stApp {
    background:
      radial-gradient(1200px 400px at 8% -5%, rgba(39,174,96,0.15), transparent 60%),
      radial-gradient(1000px 380px at 95% -8%, rgba(52,152,219,0.18), transparent 60%),
      linear-gradient(180deg, #0d1117 0%, #0b1220 100%);
}

.block-container {
    padding-top: 1.1rem;
    padding-bottom: 1.8rem;
    max-width: 100%;
    padding-right: 390px;
}

.hero {
    background:
      radial-gradient(260px 120px at 8% 5%, rgba(0,224,255,0.20), transparent 70%),
      radial-gradient(260px 120px at 92% 18%, rgba(0,255,153,0.16), transparent 70%),
      linear-gradient(120deg, rgba(16,24,39,0.88), rgba(10,30,54,0.88));
    border: 1px solid rgba(255,255,255,0.15);
    border-radius: 18px;
    padding: 18px;
    margin-bottom: 14px;
}

.brand-row {
    display: flex;
    align-items: center;
    gap: 12px;
}

.logo-orb {
    width: 42px;
    height: 42px;
    border-radius: 12px;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 0.92rem;
    font-weight: 700;
    letter-spacing: 0.4px;
    color: #03131a;
    background: linear-gradient(145deg, #38bdf8, #22d3ee 45%, #34d399);
    box-shadow: 0 0 0 2px rgba(255,255,255,0.16), 0 10px 22px rgba(34,211,238,0.35);
}

.brand-name {
    font-size: 1.85rem;
    font-weight: 700;
    letter-spacing: 0.3px;
    line-height: 1;
    background: linear-gradient(90deg, #f8fafc 0%, #89f5ff 45%, #5bf7c8 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
}

.hero-sub {
    margin-top: 8px;
    font-size: 0.96rem;
    font-weight: 600;
    letter-spacing: 0.2px;
    color: #dbeafe;
    text-shadow: 0 0 10px rgba(56,189,248,0.25);
}

.sticker-row {
    margin-top: 12px;
    display: flex;
    flex-wrap: wrap;
    gap: 8px;
}

.sticker {
    display: inline-flex;
    align-items: center;
    gap: 6px;
    padding: 6px 12px;
    border-radius: 999px;
    background: linear-gradient(120deg, rgba(15,23,42,0.8), rgba(30,41,59,0.8));
    border: 1px solid rgba(125,211,252,0.35);
    box-shadow: inset 0 0 0 1px rgba(255,255,255,0.06), 0 4px 12px rgba(0,0,0,0.25);
    color: #e2e8f0;
    font-size: 0.79rem;
    font-weight: 600;
    transition: transform 140ms ease, box-shadow 140ms ease, border-color 140ms ease;
}

.sticker:hover {
    transform: translateY(-1px);
    border-color: rgba(45,212,191,0.7);
    box-shadow: 0 0 0 1px rgba(45,212,191,0.35), 0 8px 16px rgba(45,212,191,0.18);
}

.card {
    background: linear-gradient(180deg, rgba(255,255,255,0.05), rgba(255,255,255,0.02));
    border: 1px solid rgba(255,255,255,0.14);
    border-radius: 16px;
    padding: 16px;
    margin-bottom: 14px;
    box-shadow: 0 10px 28px rgba(0,0,0,0.2);
}

.section-head {
    font-size: 1.1rem;
    font-weight: 700;
    margin-bottom: 8px;
}

.muted {
    font-size: 0.86rem;
    opacity: 0.83;
}

hr.slim {
    border: none;
    border-top: 1px solid rgba(255,255,255,0.14);
    margin: 10px 0 12px 0;
}

.log-pill {
    display: inline-block;
    border: 1px solid rgba(255,255,255,0.2);
    border-radius: 999px;
    padding: 2px 8px;
    font-size: 0.74rem;
    margin-right: 6px;
    margin-bottom: 6px;
    opacity: 0.92;
}

[data-testid="stExpander"] {
    border: 1px solid rgba(255,255,255,0.14);
    border-radius: 12px;
    background: rgba(255,255,255,0.02);
}

[data-testid="stExpander"] details summary {
    font-weight: 600;
    font-size: 0.93rem;
}

.stButton > button {
    border-radius: 10px;
    border: 1px solid rgba(255,255,255,0.16);
    font-weight: 600;
}

[data-testid="stDownloadButton"] button {
    border-radius: 10px;
    border: 1px solid rgba(255,255,255,0.16);
    font-weight: 600;
}

.think-search-note {
    font-size: 0.78rem;
    opacity: 0.75;
    margin-top: 2px;
}

.composer-actions [data-testid="stButton"] button {
    height: 52px;
}

/* Pin the top-level right panel while page scrolls */
section.main > div > div > div > div[data-testid="stHorizontalBlock"] > div[data-testid="column"]:nth-child(2) {
    position: fixed;
    top: 0;
    right: 0;
    width: 370px;
    height: 100vh;
    overflow-y: auto;
    padding: 12px 14px;
    background: linear-gradient(180deg, rgba(8,15,30,0.95), rgba(5,12,24,0.96));
    border-left: 1px solid rgba(255,255,255,0.12);
    z-index: 999;
}

section.main > div > div > div > div[data-testid="stHorizontalBlock"] > div[data-testid="column"]:nth-child(2) > div {
    padding-top: 4px;
}

section.main > div > div > div > div[data-testid="stHorizontalBlock"] > div[data-testid="column"]:nth-child(2)::-webkit-scrollbar {
    width: 8px;
}

section.main > div > div > div > div[data-testid="stHorizontalBlock"] > div[data-testid="column"]:nth-child(2)::-webkit-scrollbar-thumb {
    background: rgba(148, 163, 184, 0.35);
    border-radius: 999px;
}

@media (max-width: 1100px) {
    .block-container {
        padding-right: 1rem;
        max-width: 100%;
    }
    section.main > div > div > div > div[data-testid="stHorizontalBlock"] > div[data-testid="column"]:nth-child(2) {
        position: static;
        width: auto;
        height: auto;
        overflow: visible;
        padding: 0;
        border-left: none;
        background: transparent;
    }
}

/* Premium SaaS polish */
.stApp::before {
    content: "";
    position: fixed;
    inset: -20% -10% auto -10%;
    height: 55vh;
    background:
      radial-gradient(40% 55% at 20% 20%, rgba(67, 233, 255, 0.20), transparent 70%),
      radial-gradient(40% 55% at 80% 10%, rgba(34, 211, 238, 0.22), transparent 70%),
      radial-gradient(45% 50% at 50% 30%, rgba(56, 189, 248, 0.15), transparent 75%);
    animation: driftGlow 18s ease-in-out infinite alternate;
    pointer-events: none;
    z-index: 0;
}

@keyframes driftGlow {
    0% { transform: translateX(-2%) translateY(0); opacity: 0.8; }
    100% { transform: translateX(2%) translateY(1.5%); opacity: 1; }
}

.hero {
    backdrop-filter: blur(8px);
    box-shadow: 0 14px 34px rgba(1, 10, 25, 0.48), inset 0 0 24px rgba(45, 212, 191, 0.08);
    animation: heroSweep 16s linear infinite;
    background-size: 170% 170%;
}

@keyframes heroSweep {
    0% { background-position: 0% 40%; }
    100% { background-position: 100% 60%; }
}

.section-head {
    font-size: 1.42rem;
    letter-spacing: 0.2px;
    margin-bottom: 12px;
}

.left-panel, .final-panel, .quote-panel, .log-panel {
    background: linear-gradient(170deg, rgba(13, 22, 40, 0.62), rgba(7, 14, 29, 0.54));
    border: 1px solid rgba(125, 211, 252, 0.20);
    border-radius: 18px;
    padding: 14px 14px 12px 14px;
    backdrop-filter: blur(8px);
    box-shadow: 0 10px 30px rgba(0, 0, 0, 0.32), inset 0 0 28px rgba(56, 189, 248, 0.07);
}

[data-testid="stVerticalBlockBorderWrapper"] {
    border: 1px solid rgba(125, 211, 252, 0.20) !important;
    border-radius: 16px !important;
    background: linear-gradient(170deg, rgba(13, 22, 40, 0.62), rgba(7, 14, 29, 0.54)) !important;
    backdrop-filter: blur(8px);
    box-shadow: 0 10px 30px rgba(0, 0, 0, 0.32), inset 0 0 24px rgba(56, 189, 248, 0.06);
}

.quote-panel {
    margin: 10px 0 16px 0;
}

.run-divider {
    position: relative;
}

.run-divider::after {
    content: "";
    position: absolute;
    top: 0;
    right: -16px;
    width: 1px;
    height: 100%;
    background: linear-gradient(180deg, transparent, rgba(56,189,248,0.6), transparent);
    box-shadow: 0 0 12px rgba(56,189,248,0.45);
}

.stButton > button, [data-testid="stDownloadButton"] button {
    border-radius: 14px !important;
    border: 1px solid rgba(125,211,252,0.35) !important;
    background: linear-gradient(145deg, rgba(12, 22, 36, 0.85), rgba(20, 35, 58, 0.70)) !important;
    transition: all 170ms ease !important;
    box-shadow: 0 4px 14px rgba(5, 10, 20, 0.35);
}

.stButton > button:hover, [data-testid="stDownloadButton"] button:hover {
    border-color: rgba(45,212,191,0.72) !important;
    box-shadow: 0 0 0 1px rgba(45,212,191,0.22), 0 0 18px rgba(34,211,238,0.22);
    transform: translateY(-1px);
}

[data-testid="stExpander"] {
    border: 1px solid rgba(125, 211, 252, 0.22);
    background: linear-gradient(160deg, rgba(15,23,42,0.56), rgba(15,23,42,0.36));
    transition: all 180ms ease;
}

[data-testid="stExpander"] details > div {
    animation: expandIn 190ms ease;
}

@keyframes expandIn {
    from { opacity: 0; transform: translateY(-2px); }
    to { opacity: 1; transform: translateY(0); }
}

.status-pill {
    display: flex;
    align-items: center;
    gap: 10px;
    border: 1px solid rgba(125, 211, 252, 0.25);
    border-radius: 999px;
    padding: 6px 12px;
    width: fit-content;
    font-weight: 600;
}

.status-dot {
    width: 10px;
    height: 10px;
    border-radius: 50%;
    display: inline-block;
}

.status-idle .status-dot { background: #94a3b8; }
.status-running .status-dot { background: #22d3ee; box-shadow: 0 0 0 0 rgba(34,211,238,0.8); animation: pulseDot 1.2s infinite; }
.status-done .status-dot { background: #22c55e; box-shadow: 0 0 8px rgba(34,197,94,0.8); }

@keyframes pulseDot {
    0% { box-shadow: 0 0 0 0 rgba(34,211,238,0.75); }
    70% { box-shadow: 0 0 0 10px rgba(34,211,238,0); }
    100% { box-shadow: 0 0 0 0 rgba(34,211,238,0); }
}

.run-loader {
    margin: 8px 0 12px 0;
    padding: 8px 12px;
    border: 1px solid rgba(125, 211, 252, 0.20);
    border-radius: 12px;
    background: rgba(15, 23, 42, 0.45);
    font-size: 0.92rem;
}

.loader-steps span {
    opacity: 0.35;
    margin-right: 8px;
    animation: cycleStep 5s linear infinite;
}

.loader-steps span:nth-child(2) { animation-delay: 1.1s; }
.loader-steps span:nth-child(3) { animation-delay: 2.2s; }
.loader-steps span:nth-child(4) { animation-delay: 3.3s; }

@keyframes cycleStep {
    0%, 15% { opacity: 1; color: #67e8f9; }
    25%, 100% { opacity: 0.35; color: #cbd5e1; }
}

.log-tag {
    display: inline-block;
    border-radius: 999px;
    padding: 3px 10px;
    font-size: 0.78rem;
    font-weight: 700;
    margin-right: 6px;
}

.tag-plan { background: rgba(34,211,238,0.22); border: 1px solid rgba(34,211,238,0.5); color: #a5f3fc; }
.tag-search { background: rgba(59,130,246,0.24); border: 1px solid rgba(96,165,250,0.5); color: #bfdbfe; }
.tag-tool { background: rgba(251,146,60,0.24); border: 1px solid rgba(251,146,60,0.55); color: #fed7aa; }
.tag-eval { background: rgba(168,85,247,0.24); border: 1px solid rgba(192,132,252,0.55); color: #e9d5ff; }
.tag-final { background: rgba(34,197,94,0.23); border: 1px solid rgba(74,222,128,0.55); color: #bbf7d0; }
.tag-default { background: rgba(148,163,184,0.22); border: 1px solid rgba(148,163,184,0.45); color: #e2e8f0; }
</style>
""",
    unsafe_allow_html=True,
)

# ---------- Session State ----------
if "final_answer" not in st.session_state:
    st.session_state.final_answer = ""
if "thinking_log" not in st.session_state:
    st.session_state.thinking_log = []
if "agent_status" not in st.session_state:
    st.session_state.agent_status = "idle"
if "agent_phase" not in st.session_state:
    st.session_state.agent_phase = "Waiting..."


def _strip_md(text: str) -> str:
    if not text:
        return ""
    t = text.strip()
    t = re.sub(r"^\*+\s*|\s*\*+$", "", t)
    t = re.sub(r"^[-*]\s+", "", t)
    t = re.sub(r"^\d+\.\s+", "", t)
    return t.strip()


def step_tag_class(step: str) -> str:
    key = (step or "").strip().lower()
    mapping = {
        "plan_created": "tag-plan",
        "search_query": "tag-search",
        "tool_search_call": "tag-tool",
        "evaluation": "tag-eval",
        "final_answer": "tag-final",
    }
    return mapping.get(key, "tag-default")


def parse_answer_sections(answer: str) -> list[tuple[str, list[str]]]:
    headings = [
        "Answer",
        "Key points",
        "Contradictions / uncertainty",
        "Citations",
        "Confidence",
    ]
    section_map = {h.lower(): h for h in headings}
    sections = {h: [] for h in headings}
    current = "Answer"

    for raw_line in (answer or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue

        matched = False
        for h in headings:
            pattern = rf"^\*{{0,2}}{re.escape(h)}\*{{0,2}}:\s*(.*)$"
            m = re.match(pattern, line, flags=re.IGNORECASE)
            if m:
                current = section_map[h.lower()]
                tail = _strip_md(m.group(1))
                if tail:
                    sections[current].append(tail)
                matched = True
                break
        if matched:
            continue

        sections[current].append(_strip_md(line))

    ordered = []
    for h in headings:
        if sections[h]:
            ordered.append((h, sections[h]))
    if not ordered and answer:
        ordered = [("Answer", [_strip_md(x) for x in answer.splitlines() if x.strip()])]
    return ordered


def build_docx_bytes(answer: str) -> bytes | None:
    if Document is None:
        return None
    doc = Document()
    doc.add_heading("AutoResearch Final Answer", level=1)
    for heading, lines in parse_answer_sections(answer):
        doc.add_heading(heading, level=2)
        for line in lines:
            if line.startswith("http://") or line.startswith("https://"):
                doc.add_paragraph(line)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
            else:
                doc.add_paragraph(line)
    buff = io.BytesIO()
    doc.save(buff)
    return buff.getvalue()


def _wrap_text(text: str, max_chars: int = 95) -> list[str]:
    words = text.split()
    if not words:
        return [""]
    lines = []
    cur = []
    length = 0
    for w in words:
        add_len = len(w) + (1 if cur else 0)
        if length + add_len > max_chars:
            lines.append(" ".join(cur))
            cur = [w]
            length = len(w)
        else:
            cur.append(w)
            length += add_len
    if cur:
        lines.append(" ".join(cur))
    return lines


def build_pdf_bytes(answer: str) -> bytes | None:
    if canvas is None or letter is None:
        return None
    buff = io.BytesIO()
    c = canvas.Canvas(buff, pagesize=letter)
    width, height = letter
    y = height - 50

    def write_line(text: str, font: str = "Helvetica", size: int = 11, gap: int = 16):
        nonlocal y
        if y < 50:
            c.showPage()
            y = height - 50
        c.setFont(font, size)
        c.drawString(50, y, text)
        y -= gap

    write_line("AutoResearch Final Answer", font="Helvetica-Bold", size=15, gap=22)
    for heading, lines in parse_answer_sections(answer):
        write_line(heading, font="Helvetica-Bold", size=12, gap=18)
        for line in lines:
            prefix = "- " if line and not line.startswith("http") else ""
            for wrapped in _wrap_text(prefix + line, 92):
                write_line(wrapped, font="Helvetica", size=10, gap=14)
        y -= 6

    c.save()
    return buff.getvalue()


# ---------- Header ----------
st.markdown(
    """
<div class="hero">
  <div class="brand-row">
    <div class="logo-orb">AR</div>
    <div class="brand-name">AutoResearch Agent</div>
  </div>
  <div class="hero-sub">Plan → Search → Evaluate → Refine → Synthesize</div>
  <div class="sticker-row">
    <span class="sticker">⚡ Rapid Discovery</span>
    <span class="sticker">🧪 Evidence First</span>
    <span class="sticker">📌 Citation Safe</span>
    <span class="sticker">🛠️ Self-Refining Loop</span>
  </div>
</div>
""",
    unsafe_allow_html=True,
)

st.markdown(
    """
<div class="quote-panel">
  <div style="font-size: 1.02rem; font-weight: 600; line-height: 1.45;">
    "This agent does not just search faster, it thinks in loops until evidence becomes clarity."
  </div>
</div>
""",
    unsafe_allow_html=True,
)

# ---------- Two-column layout ----------
left_col, right_col = st.columns([2, 1], gap="large")

# ======================
# LEFT (2/3): Main UI
# ======================
with left_col:
    with st.container(border=True):
        st.markdown('<div class="section-head">❓ Question</div>', unsafe_allow_html=True)

        q_col, a_col = st.columns([6.3, 2.7], gap="small")
        with q_col:
            question = st.text_area(
                "Research question",
                placeholder="Ask anything...",
                height=118,
                label_visibility="collapsed",
            )
        with a_col:
            st.markdown('<div class="composer-actions">', unsafe_allow_html=True)
            b1, b2 = st.columns([1.5, 1.0], gap="small")
            with b1:
                run_btn = st.button("🚀 Run", use_container_width=True)
            with b2:
                clear_btn = st.button("🧹 Clear", use_container_width=True, help="Clear current answer and logs")
            st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<hr class='slim'/>", unsafe_allow_html=True)
        st.markdown('<div class="section-head">⚙️ Run Options</div>', unsafe_allow_html=True)
        o1, o2, o3 = st.columns(3)

        with o1:
            max_iters = st.slider(
                "Max iterations",
                min_value=1,
                max_value=6,
                value=3,
                help="Higher value explores more loops.",
            )
        with o2:
            max_results = st.slider(
                "Results per search",
                min_value=1,
                max_value=8,
                value=3,
                help="Higher value gathers more sources each loop.",
            )
        with o3:
            mode = st.selectbox(
                "Depth mode",
                ["Fast", "Balanced", "Deep"],
                index=1,
                help="Deep returns longer, richer final answers.",
            )

    if clear_btn:
        st.session_state.final_answer = ""
        st.session_state.thinking_log = []
        st.session_state.agent_status = "idle"
        st.session_state.agent_phase = "Waiting..."

    if run_btn:
        if not question.strip():
            st.error("Please enter a research question.")
        else:
            st.session_state.agent_status = "running"
            st.session_state.agent_phase = "Planning..."
            with st.spinner("Running agent..."):
                st.markdown(
                    """
<div class="run-loader">
  <div><strong>Running</strong> <span style="opacity:.8">-></span> <span id="run-phase">Planning...</span></div>
  <div class="loader-steps" style="margin-top:5px;">
    <span>Planning...</span><span>Searching...</span><span>Evaluating...</span><span>Synthesizing...</span>
  </div>
</div>
""",
                    unsafe_allow_html=True,
                )
                final_answer, thinking_log = run_agent(
                    question.strip(),
                    max_iterations=max_iters,
                    max_results=max_results,
                    mode=mode,
                )
            st.session_state.final_answer = final_answer
            st.session_state.thinking_log = thinking_log
            st.session_state.agent_status = "completed"
            st.session_state.agent_phase = "Completed"
            st.rerun()

    with st.container(border=True):
        st.markdown('<div class="section-head">✅ Final Answer</div>', unsafe_allow_html=True)
        st.markdown("<hr class='slim'/>", unsafe_allow_html=True)

        if st.session_state.final_answer:
            st.markdown(st.session_state.final_answer)
            st.markdown("<hr class='slim'/>", unsafe_allow_html=True)
            st.markdown('<div class="muted">Export Final Answer</div>', unsafe_allow_html=True)

            pdf_data = build_pdf_bytes(st.session_state.final_answer)
            docx_data = build_docx_bytes(st.session_state.final_answer)

            d1, d2, d3 = st.columns(3)
            with d1:
                st.download_button(
                    "⬇️ Download .md",
                    data=st.session_state.final_answer,
                    file_name="final_answer.md",
                    mime="text/markdown",
                    use_container_width=True,
                )
            with d2:
                st.download_button(
                    "⬇️ Download PDF",
                    data=pdf_data if pdf_data else b"",
                    file_name="final_answer.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    disabled=pdf_data is None,
                )
            with d3:
                st.download_button(
                    "⬇️ Download DOCX",
                    data=docx_data if docx_data else b"",
                    file_name="final_answer.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    disabled=docx_data is None,
                )

            if pdf_data is None or docx_data is None:
                st.caption("To enable PDF/DOCX export: `pip install reportlab python-docx`")
        else:
            st.info("Run a question to see the final answer.")

# ==========================
# RIGHT (1/3): Thinking Log
# ==========================
with right_col:
    with st.container(border=True):
        t1, t2 = st.columns([2.0, 1.2])
        with t1:
            st.markdown('<div class="section-head">🔎 Thinking Log</div>', unsafe_allow_html=True)
        with t2:
            st.download_button(
                "📥 Download Tlog",
                data=json.dumps(st.session_state.thinking_log, indent=2),
                file_name="thinking_log.json",
                mime="application/json",
                use_container_width=True,
            )

        status = st.session_state.agent_status
        if status == "running":
            st.markdown('<div class="status-pill status-running"><span class="status-dot"></span>Running</div>', unsafe_allow_html=True)
        elif status == "completed":
            st.markdown('<div class="status-pill status-done"><span class="status-dot"></span>Completed ✓</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="status-pill status-idle"><span class="status-dot"></span>Idle</div>', unsafe_allow_html=True)

        st.markdown(
            '<div class="muted">Live internal timeline: plan, queries, tool calls, evaluations, and synthesis.</div>',
            unsafe_allow_html=True,
        )
        st.markdown('<div class="think-search-note">Filter and scroll through events like a sidebar feed.</div>', unsafe_allow_html=True)
        st.markdown("<hr class='slim'/>", unsafe_allow_html=True)

        logs = st.session_state.thinking_log
        if not logs:
            st.info("Run a question to populate the timeline.")
        else:
            step_options = sorted({str(x.get("step", "event")) for x in logs})
            f1, f2 = st.columns([1.6, 1.0])
            with f1:
                log_search = st.text_input("Search events", placeholder="search text, step, data...", key="log_search")
            with f2:
                step_filter = st.selectbox("Step", ["All"] + step_options, index=0, key="log_step_filter")
            newest_first = st.toggle("Newest first", value=True, key="log_sort_desc")

            filtered = []
            for item in logs:
                step_val = str(item.get("step", "event"))
                if step_filter != "All" and step_val != step_filter:
                    continue
                if log_search:
                    blob = json.dumps(item, ensure_ascii=False).lower()
                    if log_search.lower() not in blob:
                        continue
                filtered.append(item)

            ordered_logs = list(reversed(filtered)) if newest_first else filtered
            st.caption(f"Showing {len(ordered_logs)} / {len(logs)} events")

            with st.container(height=560, border=False):
                for i, item in enumerate(ordered_logs, start=1):
                    ts = item.get("time", "")
                    step = item.get("step", "event")
                    data = item.get("data", None)

                    expander_label = f"{i:02d} • {step}"
                    with st.expander(expander_label, expanded=False):
                        st.markdown(
                            f"<span class='log-pill'>⏱ {ts}</span><span class='log-tag {step_tag_class(step)}'>{step}</span>",
                            unsafe_allow_html=True,
                        )
                        if data is None:
                            st.write("-")
                        elif isinstance(data, (dict, list)):
                            st.code(json.dumps(data, indent=2, ensure_ascii=False), language="json")
                        else:
                            st.markdown(str(data))
